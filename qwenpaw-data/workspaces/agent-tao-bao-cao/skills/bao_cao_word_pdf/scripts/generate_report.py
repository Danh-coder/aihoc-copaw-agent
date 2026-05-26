#!/usr/bin/env python3
"""Generate a report from a Word template, then export to PDF."""

from __future__ import annotations

import argparse
import datetime as dt
import json
import logging
import re
import shutil
import subprocess
import sys
from copy import deepcopy
from pathlib import Path
from typing import Any

import requests
from docx import Document

RANDOM_USER_API = "https://randomuser.me/api/"
LOGGER_NAME = "qwenpaw.skill.bao_cao_word_pdf.generate_report"
logger = logging.getLogger(LOGGER_NAME)
ROW_PLACEHOLDER_RE = re.compile(r"\{\{row\d+_(name|email|phone|nationality|note)\}\}")
SUPPORTED_NAT_CODES = {
    "AU",
    "BR",
    "CA",
    "CH",
    "DE",
    "DK",
    "ES",
    "FI",
    "FR",
    "GB",
    "IE",
    "IN",
    "IR",
    "MX",
    "NL",
    "NO",
    "NZ",
    "RS",
    "TR",
    "UA",
    "US",
}


def _truncate(value: Any, max_len: int = 1200) -> str:
    text = str(value)
    if len(text) <= max_len:
        return text
    return text[:max_len] + "...<truncated>"


def _setup_logging() -> None:
    logging.basicConfig(
        level=logging.INFO,
        format=(
            "%(asctime)s | %(levelname)s | "
            "%(name)s | %(message)s"
        ),
        stream=sys.stderr,
        force=True,
    )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Fill a Word report template from Random User API and export PDF.",
    )
    parser.add_argument("--template", required=True, help="Path to .docx template")
    parser.add_argument("--output-dir", required=True, help="Output directory")

    parser.add_argument("--report-title", required=True)
    parser.add_argument("--report-period", required=True)
    parser.add_argument("--generated-by", required=True)
    parser.add_argument("--generated-date", default=dt.date.today().isoformat())
    parser.add_argument("--summary-text", required=True)

    parser.add_argument("--results", type=int, required=True)
    parser.add_argument("--nat", default="")
    parser.add_argument("--gender", choices=["male", "female", ""], default="")
    parser.add_argument("--seed", default="")
    parser.add_argument("--max-template-rows", type=int, default=5)

    args = parser.parse_args()
    args.nat = _validate_nat(args.nat)
    return args


def _validate_nat(nat: str) -> str:
    raw = (nat or "").strip()
    if not raw:
        return ""

    # RandomUser supports comma-separated nationality codes.
    tokens = [piece.strip().upper() for piece in raw.split(",") if piece.strip()]
    invalid = [code for code in tokens if code not in SUPPORTED_NAT_CODES]
    if invalid:
        supported = ",".join(sorted(SUPPORTED_NAT_CODES))
        raise ValueError(
            "Unsupported --nat code(s): "
            f"{','.join(invalid)}. RandomUser supports: {supported}."
        )
    return ",".join(tokens)


def fetch_random_users(
    *,
    results: int,
    nat: str,
    gender: str,
    seed: str,
) -> list[dict[str, Any]]:
    if results <= 0:
        raise ValueError("results must be > 0")

    params: dict[str, str | int] = {"results": results}
    if nat.strip():
        params["nat"] = nat.strip()
    if gender.strip():
        params["gender"] = gender.strip()
    if seed.strip():
        params["seed"] = seed.strip()

    response = requests.get(
        RANDOM_USER_API,
        params=params,
        timeout=20,
    )
    response.raise_for_status()

    payload = response.json()
    users = payload.get("results")
    if not isinstance(users, list) or not users:
        raise RuntimeError("Random User API returned empty results")
    return users


def _safe_get(dct: dict[str, Any], *keys: str, default: str = "") -> str:
    current: Any = dct
    for key in keys:
        if not isinstance(current, dict):
            return default
        current = current.get(key)
    if current is None:
        return default
    return str(current)


def build_placeholders(
    *,
    report_title: str,
    report_period: str,
    generated_by: str,
    generated_date: str,
    summary_text: str,
    users: list[dict[str, Any]],
    max_template_rows: int,
) -> dict[str, str]:
    placeholders = {
        "{{report_title}}": report_title,
        "{{report_period}}": report_period,
        "{{generated_by}}": generated_by,
        "{{generated_date}}": generated_date,
        "{{summary_text}}": summary_text,
    }

    for idx in range(1, max_template_rows + 1):
        user = users[idx - 1] if idx - 1 < len(users) else {}
        first = _safe_get(user, "name", "first")
        last = _safe_get(user, "name", "last")
        full_name = (first + " " + last).strip()

        placeholders[f"{{{{row{idx}_name}}}}"] = full_name
        placeholders[f"{{{{row{idx}_email}}}}"] = _safe_get(user, "email")
        placeholders[f"{{{{row{idx}_phone}}}}"] = _safe_get(user, "phone")
        placeholders[f"{{{{row{idx}_nationality}}}}"] = _safe_get(user, "nat")

        city = _safe_get(user, "location", "city")
        country = _safe_get(user, "location", "country")
        note = ", ".join([piece for piece in [city, country] if piece]).strip()
        placeholders[f"{{{{row{idx}_note}}}}"] = note

    return placeholders


def _build_user_values(user: dict[str, Any]) -> dict[str, str]:
    first = _safe_get(user, "name", "first")
    last = _safe_get(user, "name", "last")
    full_name = (first + " " + last).strip()
    city = _safe_get(user, "location", "city")
    country = _safe_get(user, "location", "country")
    note = ", ".join([piece for piece in [city, country] if piece]).strip()

    return {
        "name": full_name,
        "email": _safe_get(user, "email"),
        "phone": _safe_get(user, "phone"),
        "nationality": _safe_get(user, "nat"),
        "note": note,
    }


def _replace_row_tokens(text: str, user_values: dict[str, str]) -> str:
    def _sub(match: re.Match[str]) -> str:
        key = match.group(1)
        return user_values.get(key, "")

    return ROW_PLACEHOLDER_RE.sub(_sub, text)


def _replace_text(text: str, placeholders: dict[str, str]) -> str:
    replaced = text
    for src, target in placeholders.items():
        replaced = replaced.replace(src, target)
    return replaced


def _replace_in_paragraph(paragraph: Any, placeholders: dict[str, str]) -> None:
    original = paragraph.text
    replaced = _replace_text(original, placeholders)
    if replaced == original:
        return

    if not paragraph.runs:
        paragraph.add_run(replaced)
        return

    paragraph.runs[0].text = replaced
    for run in paragraph.runs[1:]:
        run.text = ""


def _replace_row_tokens_in_paragraph(paragraph: Any, user_values: dict[str, str]) -> None:
    original = paragraph.text
    replaced = _replace_row_tokens(original, user_values)
    if replaced == original:
        return

    if not paragraph.runs:
        paragraph.add_run(replaced)
        return

    paragraph.runs[0].text = replaced
    for run in paragraph.runs[1:]:
        run.text = ""


def _row_has_user_placeholders(row: Any) -> bool:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            if ROW_PLACEHOLDER_RE.search(paragraph.text or ""):
                return True
    return False


def _fill_table_user_rows(table: Any, users: list[dict[str, Any]]) -> bool:
    placeholder_indices: list[int] = []
    for idx, row in enumerate(table.rows):
        if _row_has_user_placeholders(row):
            placeholder_indices.append(idx)

    if not placeholder_indices:
        return False

    template_row = table.rows[placeholder_indices[0]]
    existing_rows = [table.rows[idx] for idx in placeholder_indices]

    for idx, user in enumerate(users):
        if idx < len(existing_rows):
            target_row = existing_rows[idx]
        else:
            table._tbl.append(deepcopy(template_row._tr))
            target_row = table.rows[-1]

        user_values = _build_user_values(user)
        for cell in target_row.cells:
            for paragraph in cell.paragraphs:
                _replace_row_tokens_in_paragraph(paragraph, user_values)

    # Remove unused placeholder rows if there are fewer users than template rows.
    if len(users) < len(existing_rows):
        for row in reversed(existing_rows[len(users) :]):
            table._tbl.remove(row._tr)

    return True


def apply_placeholders(
    template_path: Path,
    output_docx: Path,
    placeholders: dict[str, str],
    users: list[dict[str, Any]],
) -> None:
    doc = Document(str(template_path))

    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, placeholders)

    for table in doc.tables:
        _fill_table_user_rows(table, users)
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, placeholders)

    doc.save(str(output_docx))


def _try_docx2pdf(input_docx: Path, output_pdf: Path) -> tuple[bool, str]:
    try:
        from docx2pdf import convert  # type: ignore
    except Exception as exc:
        return False, f"docx2pdf not available: {exc}"

    try:
        convert(str(input_docx), str(output_pdf))
    except Exception as exc:
        return False, f"docx2pdf convert failed: {exc}"

    if output_pdf.exists() and output_pdf.stat().st_size > 0:
        return True, "converted by docx2pdf"
    return False, "docx2pdf did not produce output"


def _try_soffice(input_docx: Path, output_pdf: Path) -> tuple[bool, str]:
    outdir = output_pdf.parent
    expected_pdf = outdir / f"{input_docx.stem}.pdf"

    commands = []
    if shutil.which("soffice"):
        commands.append("soffice")
    if shutil.which("libreoffice"):
        commands.append("libreoffice")

    if not commands:
        return False, "LibreOffice executable not found in PATH"

    errors: list[str] = []
    for cmd in commands:
        proc = subprocess.run(
            [
                cmd,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(outdir),
                str(input_docx),
            ],
            capture_output=True,
            text=True,
            check=False,
        )
        if proc.returncode != 0:
            errors.append(f"{cmd} failed: {proc.stderr.strip()}")
            continue

        if expected_pdf.exists() and expected_pdf.stat().st_size > 0:
            if expected_pdf != output_pdf:
                expected_pdf.replace(output_pdf)
            return True, f"converted by {cmd}"

        errors.append(f"{cmd} ran but output PDF missing")

    return False, " | ".join(errors)


def _try_reportlab_fallback(
    input_docx: Path,
    output_pdf: Path,
) -> tuple[bool, str]:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
    except Exception as exc:
        return False, f"reportlab not available: {exc}"

    try:
        doc = Document(str(input_docx))
        c = canvas.Canvas(str(output_pdf), pagesize=A4)
        width, height = A4
        x = 40
        y = height - 40
        line_height = 16

        def write_line(text: str) -> None:
            nonlocal y
            clean = (text or "").replace("\n", " ").strip()
            if not clean:
                clean = " "
            if y <= 40:
                c.showPage()
                y = height - 40
            c.drawString(x, y, clean[:130])
            y -= line_height

        for paragraph in doc.paragraphs:
            write_line(paragraph.text)

        for table in doc.tables:
            write_line("-")
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                write_line(row_text)
            write_line("-")

        c.save()
    except Exception as exc:
        return False, f"reportlab fallback failed: {exc}"

    if output_pdf.exists() and output_pdf.stat().st_size > 0:
        return True, "converted by reportlab fallback"
    return False, "reportlab fallback did not produce output"


def convert_docx_to_pdf(input_docx: Path, output_pdf: Path) -> str:
    ok, msg = _try_docx2pdf(input_docx, output_pdf)
    if ok:
        return msg

    ok, msg2 = _try_soffice(input_docx, output_pdf)
    if ok:
        return msg2

    ok, msg3 = _try_reportlab_fallback(input_docx, output_pdf)
    if ok:
        return msg3

    raise RuntimeError(
        "PDF conversion failed. "
        f"docx2pdf: {msg}. "
        f"soffice: {msg2}. "
        f"reportlab: {msg3}."
    )


def sanitize_filename(name: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", name.strip())
    return cleaned.strip("_") or "bao_cao"


def main() -> None:
    _setup_logging()
    args = parse_args()

    logger.info(
        "report_script_start status=started template=%s output_dir=%s "
        "results=%s",
        args.template,
        args.output_dir,
        args.results,
    )

    template_path = Path(args.template).resolve()
    output_dir = Path(args.output_dir).resolve()

    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    output_dir.mkdir(parents=True, exist_ok=True)

    users = fetch_random_users(
        results=args.results,
        nat=args.nat,
        gender=args.gender,
        seed=args.seed,
    )
    logger.info(
        "report_script_users_fetched status=ok count=%s",
        len(users),
    )

    placeholders = build_placeholders(
        report_title=args.report_title,
        report_period=args.report_period,
        generated_by=args.generated_by,
        generated_date=args.generated_date,
        summary_text=args.summary_text,
        users=users,
        max_template_rows=args.max_template_rows,
    )

    timestamp = dt.datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = sanitize_filename(args.report_title)
    output_docx = output_dir / f"{base_name}_{timestamp}.docx"
    output_pdf = output_dir / f"{base_name}_{timestamp}.pdf"

    apply_placeholders(template_path, output_docx, placeholders, users)
    logger.info(
        "report_script_docx_written status=ok path=%s",
        output_docx,
    )

    converter = convert_docx_to_pdf(output_docx, output_pdf)
    logger.info(
        "report_script_pdf_written status=ok path=%s converter=%s",
        output_pdf,
        converter,
    )

    result = {
        "status": "ok",
        "template_path": str(template_path),
        "output_docx": str(output_docx),
        "pdf_path": str(output_pdf),
        "converter": converter,
        "users_count": len(users),
    }
    logger.info(
        "report_script_finish status=success output_docx=%s output_pdf=%s "
        "users_count=%s output_json=%s",
        output_docx,
        output_pdf,
        len(users),
        _truncate(json.dumps(result, ensure_ascii=False)),
    )
    print(json.dumps(result, ensure_ascii=False))


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        error_text = _truncate(exc)
        failure_result = {
            "status": "failed",
            "error": error_text,
            "error_type": type(exc).__name__,
        }
        try:
            print(json.dumps(failure_result, ensure_ascii=False))
        except Exception:
            # Last-resort output so caller can still detect failure.
            print('{"status":"failed","error":"unknown"}')
        logger.exception(
            "report_script_finish status=failed error=%s",
            error_text,
        )
        sys.exit(1)
